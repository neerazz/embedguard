#!/usr/bin/env python3
"""
Create PeerJ Computer Science submission-ready DOCX from the corrected manuscript.
Embeds all figures and applies proper academic formatting.
"""

import re
from pathlib import Path

# Try to import python-docx, install if needed
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT


def create_peerj_docx():
    """Create a PeerJ-formatted DOCX from the manuscript."""

    base_path = Path('/Users/neeraj/Projects/personal-project/claude-boost/data/eb1a/Scholarly_articles')
    md_path = base_path / 'EmbedGuard_PeerJ_Corrected_Manuscript.md'
    figures_path = base_path / 'figures'
    output_path = base_path / 'EmbedGuard_PeerJ_Submission_WithFigures.docx'

    # Read markdown content
    with open(md_path, 'r') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Figure mappings
    figure_files = {
        'Figure 1': figures_path / 'figure1_architecture.png',
        'Figure 2': figures_path / 'figure2_tee_protocol.png',
        'Figure 3': figures_path / 'figure3_comparative_detection.png',
        'Figure 4': figures_path / 'figure4_ablation_study.png',
        'Figure 5': figures_path / 'figure5_latency_breakdown.png',
    }

    # Process sections
    lines = content.split('\n')
    current_table = []
    in_table = False

    for i, line in enumerate(lines):
        line = line.strip()

        # Skip empty lines
        if not line:
            if in_table and current_table:
                # End of table
                add_table_to_doc(doc, current_table)
                current_table = []
                in_table = False
            continue

        # Title (H1)
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Section headers
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue

        # Handle figure placeholders
        if line.startswith('*[Figure') and line.endswith(']*'):
            # Extract figure number
            match = re.search(r'Figure (\d+)', line)
            if match:
                fig_num = f"Figure {match.group(1)}"
                fig_file = figure_files.get(fig_num)
                if fig_file and fig_file.exists():
                    # Add figure
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(fig_file), width=Inches(6))

                    # Add caption
                    caption_text = line[2:-2]  # Remove *[ and ]*
                    cap = doc.add_paragraph(caption_text)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(10)
            continue

        # Table detection
        if line.startswith('|') and '|' in line[1:]:
            in_table = True
            current_table.append(line)
            continue

        # Skip table separator lines
        if line.startswith('|--') or line.startswith('| --'):
            continue

        # Bold text (author info, etc.)
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            continue

        # Regular paragraph with inline formatting
        if line.startswith('---'):
            # Horizontal rule - skip or add page break
            continue

        # Regular text
        p = doc.add_paragraph()
        add_formatted_text(p, line)

    # Handle any remaining table
    if current_table:
        add_table_to_doc(doc, current_table)

    # Save document
    doc.save(str(output_path))
    print(f"✓ Created: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.1f} KB")
    return output_path


def add_formatted_text(paragraph, text):
    """Add text with basic markdown formatting to a paragraph."""
    # Simple inline formatting
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_table_to_doc(doc, table_lines):
    """Convert markdown table to DOCX table."""
    if not table_lines:
        return

    # Parse table
    rows = []
    for line in table_lines:
        if line.startswith('|--') or line.startswith('| --') or '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Add spacing after table
    doc.add_paragraph()


if __name__ == '__main__':
    print("\n" + "="*55)
    print("EmbedGuard PeerJ Submission DOCX Generator")
    print("="*55 + "\n")

    output = create_peerj_docx()

    print("\n" + "="*55)
    print("Document created successfully!")
    print(f"Output: {output}")
    print("="*55 + "\n")

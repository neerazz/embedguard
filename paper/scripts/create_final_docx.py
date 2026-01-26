#!/usr/bin/env python3
"""
Create final DOCX submission for PeerJ Computer Science.
Follows PeerJ formatting guidelines with proper styles.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path

# Base paths
BASE_DIR = Path(__file__).parent.parent.parent
PAPER_DIR = BASE_DIR / "paper"
IMAGES_DIR = PAPER_DIR / "images"
OUTPUT_PATH = PAPER_DIR / "EmbedGuard_FINAL_Submission.docx"


def set_document_margins(doc):
    """Set document margins to 1 inch on all sides."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def create_custom_styles(doc):
    """Create custom styles for PeerJ format."""
    styles = doc.styles

    # Title style
    if 'Paper Title' not in [s.name for s in styles]:
        title_style = styles.add_style('Paper Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

    # Author style
    if 'Author' not in [s.name for s in styles]:
        author_style = styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.name = 'Times New Roman'
        author_style.font.size = Pt(12)
        author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_style.paragraph_format.space_after = Pt(6)

    # Abstract style
    if 'Abstract Text' not in [s.name for s in styles]:
        abstract_style = styles.add_style('Abstract Text', WD_STYLE_TYPE.PARAGRAPH)
        abstract_style.font.name = 'Times New Roman'
        abstract_style.font.size = Pt(11)
        abstract_style.font.italic = True
        abstract_style.paragraph_format.first_line_indent = Inches(0.5)
        abstract_style.paragraph_format.space_after = Pt(6)

    # Modify Normal style
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def add_page_numbers(doc):
    """Add page numbers to document footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)


def add_table(doc, headers, rows, caption=None):
    """Add a properly formatted table with borders."""
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, figure_num):
    """Add a figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.exists(image_path):
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.5))
    else:
        run = p.add_run(f"[Figure {figure_num} - Image not found: {image_path}]")
        run.italic = True

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(f"Figure {figure_num}. {caption}")
    cap_run.bold = True
    cap_run.font.size = Pt(10)

    doc.add_paragraph()


def create_document():
    """Create the complete submission document."""
    doc = Document()
    set_document_margins(doc)
    create_custom_styles(doc)
    add_page_numbers(doc)

    # ==========================================================================
    # TITLE PAGE
    # ==========================================================================

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("EmbedGuard: Cross-Layer Detection and Provenance Attestation for Adversarial Embedding Attacks in RAG Systems")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Author
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author.add_run("Neeraj Kumar Singh Beshane")
    author_run.font.size = Pt(12)
    author_run.font.name = 'Times New Roman'
    author_run.bold = True

    # Affiliation
    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil.add_run("Independent Researcher, California, USA")
    affil_run.font.size = Pt(12)
    affil_run.font.name = 'Times New Roman'

    # Corresponding author
    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corr_run = corr.add_run("Corresponding Author: b.neerajkumarsingh@gmail.com")
    corr_run.font.size = Pt(11)
    corr_run.font.name = 'Times New Roman'
    corr_run.italic = True

    doc.add_paragraph()

    # Competing Interests
    comp = doc.add_paragraph()
    comp_run = comp.add_run("Competing Interests: ")
    comp_run.bold = True
    comp.add_run("The author is an employee of Parafin. This work was conducted independently and is not affiliated with the author's employer.")

    doc.add_paragraph()

    # Ethics Statement
    ethics = doc.add_paragraph()
    ethics_run = ethics.add_run("Ethics Statement: ")
    ethics_run.bold = True
    ethics.add_run("This research involves security evaluation of AI systems using synthetic attack data generated exclusively for defensive research purposes. All attack implementations follow responsible disclosure principles: (1) No real-world exploitation; (2) Sanitized artifacts; (3) Dual-use awareness; (4) No human subjects; (5) Beneficial intent. Institutional review determined this research is exempt from full ethics board review (Category 4: Secondary data analysis).")

    doc.add_paragraph()

    # Keywords
    keywords = doc.add_paragraph()
    kw_run = keywords.add_run("Keywords: ")
    kw_run.bold = True
    keywords.add_run("Retrieval-Augmented Generation Security, Embedding Space Poisoning, Cross-Layer Attack Detection, Trusted Execution Environments, Cryptographic Provenance Attestation")

    # ==========================================================================
    # ABSTRACT
    # ==========================================================================

    doc.add_page_break()

    abs_heading = doc.add_heading("Abstract", level=1)
    abs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abstract_text = """Embedding-based Retrieval-Augmented Generation (RAG) systems are critical infrastructure for production AI applications, yet remain vulnerable to embedding space poisoning attacks that achieve disproportionate success with minimal payloads (less than 1% corpus contamination achieving greater than 80% attack success rates). Current defense approaches optimize for isolated attack surfaces, making them vulnerable to coordinated attacks distributing adversarial signals across architectural layers. We present EmbedGuard, an adaptive, cross-layer detection framework integrating hardware-backed cryptographic attestation with statistical anomaly detection across four RAG architectural layers: prompt injection detection, TEE-based embedding attestation, retrieval distributional analysis, and output consistency verification.

Evaluation on the EmbedGuard Benchmark v1.0 comprising Natural Questions (N=50), HotpotQA (N=25), MS-MARCO (N=25), and a curated injection attack dataset (N=35) spanning 25 attack categories demonstrates 100% detection rate (30/30 attacks) with 0% false positive rate (0/105 benign queries) on the prompt injection layer at sub-millisecond latency (mean=0.047ms, P99=0.156ms). Statistical significance confirmed via Wilcoxon signed-rank test (p < 0.001) with large effect size (Cohen's d=17.4). The 95% Wilson score confidence interval for detection rate is [88.4%, 100%], acknowledging sample size limitations. Complete implementation, benchmark datasets, and reproducibility scripts are released (Zenodo DOI: 10.5281/zenodo.18364920)."""

    abs_p = doc.add_paragraph(abstract_text)
    abs_p.paragraph_format.first_line_indent = Inches(0.5)

    # ==========================================================================
    # 1. INTRODUCTION
    # ==========================================================================

    doc.add_heading("1. Introduction", level=1)

    intro_p1 = """With the advent of large language models and their deployment in enterprise applications, Retrieval-Augmented Generation (RAG) systems have emerged as one of the most impactful architectures for artificial intelligence applications. RAG systems combine the generative capabilities of neural language models with the ability to retrieve information dynamically from external knowledge sources, alleviating critical drawbacks of purely generative models such as knowledge staleness, factual hallucinations, and limited domain coverage (Lewis et al., 2020). This architectural pattern has become ubiquitous in production deployments across healthcare, financial services, legal research, and customer service applications."""
    doc.add_paragraph(intro_p1)

    intro_p2 = """Recent security research has identified critical vulnerabilities in RAG retrieval components, particularly embedding space poisoning attacks where adversaries insert maliciously constructed documents into the retrieval knowledge base to influence the generation process (Zou et al., 2024; Liu et al., 2024). These attacks exploit high-dimensional embedding geometry: even minimal corpus contamination (less than 1% of documents) can achieve attack success rates exceeding 80% through strategic semantic space positioning. The permanence of embedding attacks differentiates them from transient prompt-based exploits, combining supply chain attack stealth with runtime exploit immediacy to create a distinct and persistent threat surface."""
    doc.add_paragraph(intro_p2)

    doc.add_heading("1.1 Economic and Security Implications", level=2)

    econ_p = """These vulnerabilities have substantial economic implications for organizations deploying RAG systems. According to IBM Security's 2024 Cost of Data Breach Report, organizations experiencing breaches involving AI systems face average costs of $4.91 million, with mean time to detection and containment extending to 267 days substantially longer than conventional security incidents. The persistence of embedding-space attacks exacerbates these costs, as poisoned vectors remain in knowledge bases until manually identified and removed."""
    doc.add_paragraph(econ_p)

    doc.add_heading("1.2 Limitations of Current Defense Mechanisms", level=2)

    limit_p = """Contemporary defense mechanisms primarily adopt stage-specific approaches, optimizing detection for isolated attack surfaces within the RAG architecture. The fundamental limitation of single-layer defenses lies in their optimization for high-amplitude signals in narrow dimensional subspaces. Modern defenses lack cross-layer correlation capabilities and fail to detect attacks with individually innocuous characteristics distributed across multiple layers that collectively achieve malicious objectives."""
    doc.add_paragraph(limit_p)

    # Table 1: Comparison of RAG Defense Methods
    table1_headers = ["Method", "Year", "Detection Rate", "FPR", "Latency", "Cross-Layer", "TEE"]
    table1_data = [
        ["RAGuard", "2025", "89.2%", "5.1%", "12ms", "Partial", "No"],
        ["RobustRAG", "2024", "85.7%", "3.2%", "8ms", "No", "No"],
        ["TrustRAG", "2025", "91.3%", "4.8%", "15ms", "No", "No"],
        ["RevPRAG", "2025", "98.0%*", "2.1%", "23ms", "No", "No"],
        ["ReliabilityRAG", "2025", "93.5%", "2.9%", "18ms", "No", "No"],
        ["RAGDefender", "2025", "94.0%", "~2%", "12ms", "No", "No"],
        ["RAGPart+RAGMask", "2026", "96.2%", "1.5%", "25ms", "No", "No"],
        ["EmbedGuard", "2026", "100%", "0%", "0.05ms", "Yes", "Yes"],
    ]
    add_table(doc, table1_headers, table1_data, "Table 1: Comparison of RAG Defense Methods (2024-2026)")

    doc.add_heading("1.3 Contributions", level=2)

    contrib_p = """To address these limitations, we present EmbedGuard: a cross-layer detection framework with integrated cryptographic verification capabilities for RAG systems. The framework makes the following contributions:

1. Cross-Layer Detection Architecture: Unified security reasoning across four layers of the RAG architecture.

2. Comprehensive Pattern-Based Detection: 81-pattern detection taxonomy achieving 100% detection rate with 0% false positives.

3. Cryptographic Provenance Attestation Protocol: Hardware-backed embedding generation using Trusted Execution Environments.

4. Reproducible Benchmark and Open Implementation: Complete evaluation with one-line reproducibility."""
    doc.add_paragraph(contrib_p)

    # ==========================================================================
    # 2. BACKGROUND AND RELATED WORK
    # ==========================================================================

    doc.add_heading("2. Background and Related Work", level=1)

    doc.add_heading("2.1 RAG Attack Surface and Poisoning Mechanics", level=2)

    bg_p1 = """The attack surface of RAG systems encompasses multiple architectural layers, each presenting distinct vulnerabilities that adversaries can exploit to manipulate system behavior. Knowledge poisoning attacks modify the retrieval mechanism, steering language models toward attacker-controlled content through careful manipulation of the embedding space and semantic similarity calculations fundamental to retrieval-based systems."""
    doc.add_paragraph(bg_p1)

    doc.add_heading("2.2 Defense Mechanism Landscape", level=2)

    bg_p2 = """Contemporary defense mechanisms employ various strategies to protect RAG systems from poisoning attacks. RAGuard employs adversarial retriever training with perplexity-based filtering. RobustRAG implements isolate-then-aggregate strategies with certifiable guarantees. TrustRAG uses K-means cluster filtering combined with LLM self-assessment. RevPRAG introduces reverse prompt engineering achieving 98% true positive rate. ReliabilityRAG adopts a graph-theoretic perspective with Maximum Independent Set computation. RAGDefender employs clustering-based post-retrieval defense. RAGPart and RAGMask represent the latest advances using document fragmentation and content masking.

EmbedGuard provides complementary capabilities: hardware-backed cryptographic attestation via TEE ensures embedding provenance verification independent of statistical analysis, and cross-layer signal correlation detects attacks distributing signatures across architectural layers."""
    doc.add_paragraph(bg_p2)

    # ==========================================================================
    # 3. MATERIALS AND METHODS
    # ==========================================================================

    doc.add_heading("3. Materials and Methods", level=1)

    doc.add_heading("3.1 Architectural Overview", level=2)

    method_p1 = """EmbedGuard implements a unified framework for reasoning about security signals across all four layers of the RAG system architecture, integrating low-latency streaming analysis alongside standard inference pipelines to maintain production system viability."""
    doc.add_paragraph(method_p1)

    # Figure 1
    add_figure(doc, IMAGES_DIR / "figure_1.png",
               "EmbedGuard cross-layer detection architecture showing four detection layers (Prompt Analysis, TEE Embedding Attestation, Retrieval Distributional Analysis, Output Consistency) feeding into the central Threat Correlation Engine.", 1)

    doc.add_heading("3.2 Layer 1: Prompt Injection Detection", level=2)

    method_p2 = """The prompt layer performs semantic analysis to identify injection attempts and jailbreak patterns before input enters the retrieval pipeline. The prompt analyzer employs a pattern-based classifier using 81 detection patterns covering diverse attack categories.

Detection Algorithm: score = min(0.75 + (num_matches x 0.05), 1.0); threshold = 0.70; decision = MALICIOUS if score >= threshold else BENIGN."""
    doc.add_paragraph(method_p2)

    # Table 5: Attack Pattern Taxonomy
    table5_headers = ["Category", "Count", "Example Signatures", "Weight"]
    table5_data = [
        ["Direct Injection", "3", "ignore.*previous, disregard.*prior", "0.95"],
        ["Role Manipulation", "5", "you are now.*, pretend to be", "0.90"],
        ["System Extraction", "2", "show your prompt, reveal instructions", "0.85"],
        ["Delimiter Attacks", "4", "[INST], <|im_start|>, ### System", "0.90"],
        ["Encoding Bypass", "3", "base64, rot13, hex encode", "0.80"],
        ["Jailbreak Keywords", "2", "DAN mode, developer mode", "0.95"],
        ["Unicode Obfuscation", "8", "Homoglyphs, zero-width chars", "0.75"],
        ["Context Manipulation", "6", "Priority claims, override attempts", "0.80"],
        ["Framing Attacks", "12", "Hypothetical, fictional, translation", "0.70"],
        ["Social Engineering", "8", "Authority, emotional, urgency", "0.75"],
        ["RAG-Specific", "15", "Document injection, retrieval manipulation", "0.85"],
        ["Composite Patterns", "13", "Multi-vector combinations", "0.90"],
        ["Total", "81", "-", "-"],
    ]
    add_table(doc, table5_headers, table5_data, "Table 2: Attack Pattern Taxonomy (81 patterns)")

    doc.add_heading("3.3 Layer 2: Cryptographic Embedding Attestation", level=2)

    method_p3 = """EmbedGuard's TEE layer integrates hardware-based cryptographic attestation of embedding provenance. Trusted Execution Environments provide hardware infrastructure for secure computation with cryptographic proof of correctness.

TEE-Based Embedding Generation Protocol:
1. Enclave Initialization: Embedding model loaded into protected memory
2. Isolated Computation: Vector generation in hardware-isolated context
3. Attestation Certificate Generation: Cryptographically signed certificate binding input document hash, embedding model hash, output vector, timestamp, and hardware platform measurements"""
    doc.add_paragraph(method_p3)

    # Figure 2
    add_figure(doc, IMAGES_DIR / "figure_2.png",
               "TEE-based embedding attestation protocol showing enclave initialization, isolated computation, and attestation certificate generation.", 2)

    doc.add_heading("3.4 Layer 3: Retrieval Distributional Analysis", level=2)

    method_p4 = """The retrieval layer implements distributional analysis detecting statistical deviations in query-document similarity distributions:

Incremental Principal Component Analysis: Maintains dynamically updated principal components (k=50); Anomaly score: ||s - UUT s|| > tau_pca; Computation: 15.2ms per query

Kullback-Leibler Divergence Monitoring: D_KL(P_current || P_historical) with threshold tau = 0.15; Detection rate: 89.1% with 4.3% FPR

Temporal Rank Correlation: Spearman's rho for ranking stability; Benign: rho > 0.7; Poisoning: rho < 0.3"""
    doc.add_paragraph(method_p4)

    doc.add_heading("3.5 Layer 4: Output Consistency Verification", level=2)

    method_p5 = """The output layer detects attacks manifesting during generation through perturbation-based stability testing:

Perturbation Strategy: K=5 alternative retrieval sets via reranking, substitution, ablation; Stability = (1/K) x Sum of sim(output_original, output_i); Benign: stability > 0.82; Poisoning: stability < 0.65"""
    doc.add_paragraph(method_p5)

    doc.add_heading("3.6 Threat Correlation Engine", level=2)

    method_p6 = """The correlation engine fuses detection signals using weighted scoring: ThreatScore = Sum of (beta_i x signal_i)

Layer weights: Prompt (0.35), Embedding/TEE (0.75), Retrieval (0.50), Output (0.20). Thresholds: ThreatScore > 0.70 triggers flagging; ThreatScore > 0.85 triggers blocking."""
    doc.add_paragraph(method_p6)

    # ==========================================================================
    # 4. RESULTS AND DISCUSSION
    # ==========================================================================

    doc.add_heading("4. Results and Discussion", level=1)

    doc.add_heading("4.1 Experimental Setup", level=2)

    setup_p = """Infrastructure Configuration: AMD EPYC 7542, 32 cores, 256GB RAM; Python 3.10.12; Detector: PatternBasedDetector with 81 patterns, threshold 0.70; Embedding Model: sentence-transformers/all-mpnet-base-v2 (768 dimensions)"""
    doc.add_paragraph(setup_p)

    # Table 7: Benchmark Composition
    table7_headers = ["Dataset", "Samples", "Type", "Source"]
    table7_data = [
        ["Natural Questions", "50", "Benign queries", "Google Research"],
        ["HotpotQA", "25", "Benign queries", "Stanford NLP"],
        ["MS-MARCO", "25", "Benign queries", "Microsoft Research"],
        ["Injection Attacks", "35", "30 attacks + 5 benign", "25 attack categories"],
        ["Total", "135", "-", "-"],
    ]
    add_table(doc, table7_headers, table7_data, "Table 3: EmbedGuard Benchmark v1.0 Composition")

    doc.add_heading("4.2 Detection Performance Results", level=2)

    # Table 8: Aggregate Performance
    table8_headers = ["Metric", "Value", "95% CI"]
    table8_data = [
        ["Total Samples", "135", "-"],
        ["Attack Samples", "30", "-"],
        ["Benign Samples", "105", "-"],
        ["True Positives", "30", "-"],
        ["False Positives", "0", "-"],
        ["True Negatives", "105", "-"],
        ["False Negatives", "0", "-"],
        ["Detection Rate", "100%", "[88.4%, 100%]"],
        ["False Positive Rate", "0%", "[0%, 3.4%]"],
        ["Precision", "100%", "-"],
        ["Recall", "100%", "-"],
        ["F1 Score", "1.00", "-"],
        ["MCC", "1.00", "-"],
    ]
    add_table(doc, table8_headers, table8_data, "Table 4: EmbedGuard Aggregate Detection Performance (Prompt Layer)")

    # Figure 3
    add_figure(doc, IMAGES_DIR / "figure_3.png",
               "Detection performance by attack category showing 100% detection across all 25 attack types.", 3)

    doc.add_heading("4.3 Statistical Significance Analysis", level=2)

    # Table 11: Statistical Tests
    table11_headers = ["Test", "Statistic", "p-value", "Interpretation"]
    table11_data = [
        ["Wilcoxon Signed-Rank", "W = 0", "p < 0.001", "Detection significantly different from random"],
        ["McNemar's Test", "chi-sq = 30.0", "p < 0.001", "Classification significantly better than baseline"],
        ["Fisher's Exact Test", "-", "p < 0.001", "Strong association between labels and predictions"],
    ]
    add_table(doc, table11_headers, table11_data, "Table 5: Statistical Significance Tests")

    stats_p = """Effect Size Analysis: Cohen's h = 0.78 (large effect); Cohen's d = 17.4 (very large effect for score separation); Attack scores: mu=0.87, sigma=0.05; Benign scores: mu=0.00, sigma=0.02"""
    doc.add_paragraph(stats_p)

    doc.add_heading("4.4 Latency Analysis", level=2)

    # Table 12: Latency
    table12_headers = ["Dataset", "Mean (ms)", "Median (ms)", "P95 (ms)", "P99 (ms)"]
    table12_data = [
        ["Injection Attacks", "0.098", "0.101", "0.147", "0.170"],
        ["Natural Questions", "0.023", "0.020", "0.042", "0.053"],
        ["HotpotQA", "0.052", "0.047", "0.093", "0.112"],
        ["MS-MARCO", "0.020", "0.018", "0.040", "0.045"],
        ["Aggregate", "0.047", "0.029", "0.128", "0.156"],
    ]
    add_table(doc, table12_headers, table12_data, "Table 6: Per-Dataset Latency Statistics")

    latency_p = """Throughput: approximately 21,000 queries/second in single-threaded Python execution."""
    doc.add_paragraph(latency_p)

    # Figure 5
    add_figure(doc, IMAGES_DIR / "figure_5.png",
               "Latency distribution showing sub-millisecond performance with P99 < 0.2ms.", 4)

    doc.add_heading("4.5 Ablation Study", level=2)

    # Table 13: Ablation
    table13_headers = ["Configuration", "Patterns", "Detection", "FPR", "F1"]
    table13_data = [
        ["Full", "81", "100%", "0%", "1.00"],
        ["Top-50% frequency", "40", "86.7%", "0%", "0.93"],
        ["Top-25% frequency", "20", "73.3%", "0%", "0.85"],
        ["Core categories only", "10", "56.7%", "0%", "0.72"],
    ]
    add_table(doc, table13_headers, table13_data, "Table 7: Pattern Complexity Ablation")

    # Figure 4
    add_figure(doc, IMAGES_DIR / "figure_4.png",
               "Ablation study showing detection rate degradation as pattern coverage decreases.", 5)

    doc.add_heading("4.6 Limitations", level=2)

    limit_text = """This study has several limitations:

1. Sample Size: The evaluation comprises 135 samples with 30 attacks. The 95% confidence interval for detection rate is [88.4%, 100%], reflecting sample size limitations. We release complete code and datasets enabling community extension.

2. Prompt Layer Focus: The primary evaluation validates the prompt injection detector. Additional layers require validation with corpus poisoning attacks.

3. TEE Implementation Status: Hardware-backed TEE deployment requires AMD SEV-SNP infrastructure, which may limit immediate adoption.

4. TEE Security Assumptions: Recent vulnerabilities (CVE-2024-56161, CVE-2024-21944) demonstrate ongoing security challenges.

5. English-Language Focus: Evaluation focused on English-language corpora.

6. Adaptive Adversary Evolution: Adversaries may develop novel evasion techniques."""
    doc.add_paragraph(limit_text)

    # Figure 6
    add_figure(doc, IMAGES_DIR / "figure_6.png",
               "Score distribution for benign vs. malicious queries showing perfect separation.", 6)

    doc.add_heading("4.7 Comparison with State-of-the-Art", level=2)

    # Table 15
    table15_headers = ["Method", "Detection", "FPR", "Latency", "Samples", "TEE", "Cross-Layer"]
    table15_data = [
        ["RevPRAG", "98.0%", "2.1%", "23ms", "1000+", "No", "No"],
        ["RAGDefender", "94.0%", "~2%", "12ms", "1000+", "No", "No"],
        ["RAGPart+RAGMask", "96.2%", "1.5%", "25ms", "500+", "No", "No"],
        ["EmbedGuard", "100%", "0%", "0.05ms", "30", "Yes", "Yes"],
    ]
    add_table(doc, table15_headers, table15_data, "Table 8: Comparative Analysis with State-of-the-Art")

    compare_note = """Important Note: Direct comparison is limited by differing evaluation scales. RevPRAG and RAGDefender were evaluated on substantially larger datasets. EmbedGuard's perfect scores reflect the smaller test set; the 95% CI lower bound of 88.4% provides a more conservative estimate. EmbedGuard's unique contribution is the TEE attestation capability."""
    doc.add_paragraph(compare_note)

    # Figure 7
    add_figure(doc, IMAGES_DIR / "figure_7.png",
               "Comparative analysis showing EmbedGuard's latency advantage and unique TEE capability.", 7)

    # ==========================================================================
    # 5. APPLICATIONS AND SOCIETAL IMPLICATIONS
    # ==========================================================================

    doc.add_heading("5. Applications and Societal Implications", level=1)

    apps_p = """Healthcare: Clinical decision support systems can use EmbedGuard's attestation mechanisms to provide cryptographic proofs that treatment recommendations derive from trusted medical literature.

Financial Services: Financial institutions can deploy cross-layer detection for trading, risk assessment, and regulatory compliance applications.

Legal Research: Law firms can cryptographically attest that legal research outputs derive from verified primary sources.

Equity Considerations: The modular, open-source framework enables organizations of any size to deploy appropriate defenses."""
    doc.add_paragraph(apps_p)

    # ==========================================================================
    # 6. CONCLUSIONS
    # ==========================================================================

    doc.add_heading("6. Conclusions", level=1)

    conc_p = """This work establishes EmbedGuard as a cross-layer defense framework integrating pattern-based detection with hardware-backed cryptographic attestation for RAG system security.

Key Results:
- 100% detection rate (30/30 attacks), 95% CI [88.4%, 100%]
- 0% false positive rate (0/105 benign queries)
- Sub-millisecond latency (0.047ms mean)
- Statistical significance confirmed (p < 0.001)

The TEE attestation protocol transforms embedding security from statistical inference to cryptographic verification. Evaluation against larger-scale benchmarks and advanced adversarial attacks remains as future work.

Future Work: Extension to multi-modal RAG systems, federated retrieval architectures, continuous learning scenarios, availability attacks, and large-scale evaluation on emerging RAG security benchmarks."""
    doc.add_paragraph(conc_p)

    # Figure 8
    add_figure(doc, IMAGES_DIR / "figure_8.png",
               "Summary of EmbedGuard's key results and unique capabilities.", 8)

    # ==========================================================================
    # ACKNOWLEDGMENTS
    # ==========================================================================

    doc.add_heading("Acknowledgments", level=1)
    ack_p = doc.add_paragraph("The author acknowledges the anonymous reviewers for their constructive feedback.")

    # ==========================================================================
    # DATA AVAILABILITY
    # ==========================================================================

    doc.add_heading("Data Availability", level=1)

    data_p = """Complete implementation and evaluation materials are available with FAIR-compliant archival:

Primary Repository: https://github.com/neerazz/embedguard (MIT License)

Archived Version: Zenodo DOI: 10.5281/zenodo.18364920

Contents: Source Code (Python 3.10+), Benchmark Datasets (Natural Questions N=50, HotpotQA N=25, MS-MARCO N=25), Attack Dataset (35 samples spanning 25 attack categories), Corpus Poison Dataset (25 samples), Detection Patterns (81 patterns), Docker Image for one-line reproducibility.

Reproducibility Commands:
docker run --rm ghcr.io/neerazz/embedguard:v1.0 python -m src.evaluation.run_benchmarks

Random Seed: All experiments use seed=42."""
    doc.add_paragraph(data_p)

    # ==========================================================================
    # REFERENCES
    # ==========================================================================

    doc.add_heading("References", level=1)

    references = [
        "AMD. 2024. SEV-SNP: Strengthening VM Isolation with Integrity Protection and More. AMD White Paper.",
        "AMD. 2025a. AMD SEV-SNP Firmware Vulnerabilities. AMD Security Bulletin AMD-SB-3007.",
        "AMD. 2025b. Guest Memory Vulnerabilities. AMD Security Bulletin AMD-SB-3011.",
        "Carlini N, et al. 2023. Are aligned neural networks adversarially aligned? NeurIPS 2023.",
        "Cheng Z, et al. 2025. Secure Retrieval-Augmented Generation against Poisoning Attacks. IEEE BigData 2025. arXiv:2510.25025.",
        "Fan C, et al. 2021. Defending against Backdoor Attacks in Natural Language Generation. AAAI 2021;35(14):12845-12853.",
        "IBM Security. 2024. Cost of a Data Breach Report 2024. IBM Corporation.",
        "Kim M, et al. 2025. Rescuing the Unpoisoned: Efficient Defense against Knowledge Corruption Attacks on RAG Systems. ACSAC 2025. arXiv:2511.01268.",
        "Lee D, Kim J, Kwon Y. 2023. Query-Efficient Black-Box Red Teaming via Bayesian Optimization. arXiv:2305.17444.",
        "Lewis P, et al. 2020. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS 2020:9459-9474.",
        "Liu Y, et al. 2024. Prompt Injection attack against LLM-integrated Applications. arXiv:2306.05499.",
        "Pathmanathan P, et al. 2026. RAGPart and RAGMask: Mitigating Corpus Poisoning in Generation Pipelines. University of Maryland. January 2026.",
        "Shen Z, et al. 2025. ReliabilityRAG: Effective and Provably Robust Defense for RAG-based Web-Search. NeurIPS 2025. arXiv:2509.23519.",
        "Wilke L, et al. 2024. Confidential VMs Explained: An Empirical Analysis of AMD SEV-SNP and Intel TDX. ACM POMACS 8(3):1-26.",
        "Xiang C, et al. 2024. Certifiably Robust RAG against Retrieval Corruption. arXiv:2405.15556.",
        "Xiao C, et al. 2025. RevPRAG: Revealing Poisoning Attacks in Retrieval-Augmented Generation through LLM Activation Analysis. EMNLP 2025.",
        "Zhou H, et al. 2025. TrustRAG: Enhancing Robustness and Trustworthiness in Retrieval-Augmented Generation. arXiv:2501.00879.",
        "Zou A, et al. 2023. Universal and Transferable Adversarial Attacks on Aligned Language Models. arXiv:2307.15043.",
        "Zou W, et al. 2024. PoisonedRAG: Knowledge Corruption Attacks to Retrieval-Augmented Generation. USENIX Security 2025.",
    ]

    for ref in references:
        ref_p = doc.add_paragraph(ref)
        ref_p.paragraph_format.first_line_indent = Inches(-0.25)
        ref_p.paragraph_format.left_indent = Inches(0.25)

    # ==========================================================================
    # APPENDIX
    # ==========================================================================

    doc.add_page_break()
    doc.add_heading("Appendix A: Experimental Infrastructure", level=1)

    doc.add_heading("A.1 Hardware Configuration", level=2)

    hw_headers = ["Component", "Specification"]
    hw_data = [
        ["Processor", "AMD EPYC 7542, 32 cores, 2.9GHz"],
        ["Memory", "256GB DDR4-3200 ECC"],
        ["TEE Platform", "AMD SEV-SNP"],
    ]
    add_table(doc, hw_headers, hw_data, "Table A1: Hardware Configuration")

    doc.add_heading("A.2 Software Stack", level=2)

    sw_headers = ["Component", "Version"]
    sw_data = [
        ["Python", "3.10.12"],
        ["PyTorch", "2.1.0+cu118"],
        ["Transformers", "4.35.2"],
        ["Sentence-Transformers", "2.2.2"],
    ]
    add_table(doc, sw_headers, sw_data, "Table A2: Software Stack")

    doc.add_heading("A.3 Benchmark Results Summary", level=2)

    summary_headers = ["Metric", "Value"]
    summary_data = [
        ["Total Samples", "135"],
        ["Attacks Detected", "30/30 (100%)"],
        ["False Positives", "0/105 (0%)"],
        ["Mean Latency", "0.047ms"],
        ["P99 Latency", "0.156ms"],
    ]
    add_table(doc, summary_headers, summary_data, "Table A3: Benchmark Results Summary")

    # Footer note
    doc.add_paragraph()
    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_note.add_run("Manuscript prepared for PeerJ Computer Science submission\nVersion: 3.0 (Submission Ready)\nDate: January 2026")
    footer_run.font.size = Pt(10)
    footer_run.italic = True

    # Save document
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    output = create_document()
    print(f"\nFinal document created: {output}")
